import io
import shutil
import subprocess
import tempfile
from abc import ABC, abstractmethod
from pathlib import Path

import docx
import openpyxl
import pdfplumber


class ExtractionError(Exception):
    pass


class BaseExtractor(ABC):
    @abstractmethod
    def extract(self, file_bytes: bytes) -> str:
        pass


class PdfExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes) -> str:
        texts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
        return "\n".join(texts)


class DocxExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes) -> str:
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)


class XlsxExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes) -> str:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws = wb.active
        rows = []
        for row in ws.iter_rows():
            values = [str(cell.value) for cell in row if cell.value is not None]
            if values:
                rows.append(" | ".join(values))
        return "\n".join(rows)


class HwpExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes) -> str:
        lo_path = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_path:
            raise ExtractionError("LibreOffice가 설치되어 있지 않습니다. HWP 변환 불가.")

        with tempfile.TemporaryDirectory() as tmpdir:
            hwp_path = Path(tmpdir) / "input.hwp"
            hwp_path.write_bytes(file_bytes)
            result = subprocess.run(
                [lo_path, "--headless", "--convert-to", "docx", "--outdir", tmpdir, str(hwp_path)],
                capture_output=True,
                timeout=60,
            )
            if result.returncode != 0:
                raise ExtractionError(f"LibreOffice 변환 실패: exit code {result.returncode}")

            docx_files = list(Path(tmpdir).glob("*.docx"))
            if not docx_files:
                raise ExtractionError("변환된 DOCX 파일을 찾을 수 없습니다.")

            return DocxExtractor().extract(docx_files[0].read_bytes())


_MIME_TO_EXTRACTOR: dict[str, type[BaseExtractor]] = {
    "application/pdf": PdfExtractor,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocxExtractor,
    "application/msword": DocxExtractor,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": XlsxExtractor,
    "application/vnd.ms-excel": XlsxExtractor,
    "application/x-hwp": HwpExtractor,
    "application/x-hwpx": HwpExtractor,
    "application/haansofthwp": HwpExtractor,
}


def get_extractor(mime_type: str) -> BaseExtractor | None:
    cls = _MIME_TO_EXTRACTOR.get(mime_type)
    return cls() if cls else None
